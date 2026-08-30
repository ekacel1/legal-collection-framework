# -*- coding: utf-8 -*-
"""Generateur DOCX pour la Software Architecture Specification LCF.

Lit un fichier source en markdown restreint et produit un .docx mis en page
selon une charte unique (arc42-like) partagee par tous les volumes.
"""
import re
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ACCENT = RGBColor(0x0B, 0x3D, 0x5B)
ACCENT2 = RGBColor(0x1B, 0x6C, 0x8F)
GREY = RGBColor(0x55, 0x55, 0x55)
FENCE = chr(96) * 3
TICK = chr(96)


_ALLOWED = ('	',)


def clean(t):
    return ''.join(c for c in t if c in _ALLOWED or (ord(c) >= 32 and ord(c) != 127))


def shade(el, color):
    sh = OxmlElement('w:shd')
    sh.set(qn('w:val'), 'clear')
    sh.set(qn('w:color'), 'auto')
    sh.set(qn('w:fill'), color)
    el.append(sh)


def borders(paragraph, color='B8C4CC', size='6', sides=('left',)):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    for s in sides:
        b = OxmlElement('w:' + s)
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:space'), '8')
        b.set(qn('w:color'), color)
        pbdr.append(b)
    pPr.append(pbdr)


def setup_styles(doc):
    st = doc.styles['Normal']
    st.font.name = 'Calibri'
    st.font.size = Pt(11)
    st.paragraph_format.space_after = Pt(8)
    st.paragraph_format.line_spacing = 1.15
    for name, size, color, bold in (
        ('Heading 1', 20, ACCENT, True),
        ('Heading 2', 15, ACCENT, True),
        ('Heading 3', 12.5, ACCENT2, True),
        ('Heading 4', 11.5, ACCENT2, True),
    ):
        s = doc.styles[name]
        s.font.name = 'Calibri'
        s.font.size = Pt(size)
        s.font.bold = bold
        s.font.color.rgb = color
        s.paragraph_format.space_before = Pt(16 if name == 'Heading 1' else 12)
        s.paragraph_format.space_after = Pt(6)
        s.paragraph_format.keep_with_next = True


def footer_text(doc, text):
    for section in doc.sections:
        p = section.footer.paragraphs[0]
        p.text = text
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.size = Pt(8)
            r.font.color.rgb = GREY


def cover(doc, meta):
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(meta.get('TITLE', ''))
    r.font.size = Pt(28)
    r.bold = True
    r.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(meta.get('SUBTITLE', ''))
    r.font.size = Pt(14)
    r.font.color.rgb = ACCENT2

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(meta.get('VOLUME', ''))
    r.font.size = Pt(18)
    r.bold = True

    for _ in range(6):
        doc.add_paragraph()
    tbl = doc.add_table(rows=0, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [
        ('Document', 'Software Architecture Specification (SAS)'),
        ('Projet', 'Legal Collection Framework - LCF'),
        ('Version', meta.get('VERSION', '0.1')),
        ('Statut', 'Architecture de reference'),
        ('Auteur', 'Chief Software Architect'),
        ('Referentiel', 'arc42 / C4 Model / ISO-IEC-IEEE 42010'),
        ('Audience', 'Architectes, developpeurs, equipes IA, mainteneurs'),
    ]
    for k, v in rows:
        c = tbl.add_row().cells
        c[0].text = k
        c[1].text = v
        for run in c[0].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
        for run in c[1].paragraphs[0].runs:
            run.font.size = Pt(10)
    doc.add_page_break()


def add_code(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Cm(0.3)
    shade(p._p.get_or_add_pPr(), 'F2F5F7')
    borders(p, color='9FB3C0', size='12', sides=('left',))
    for i, line in enumerate(lines):
        r = p.add_run(clean(line))
        r.font.name = 'Consolas'
        r.font.size = Pt(8.5)
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
        if i < len(lines) - 1:
            r.add_break()


def add_note(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    shade(p._p.get_or_add_pPr(), 'EAF2F6')
    borders(p, color='1B6C8F', size='18', sides=('left',))
    r = p.add_run(clean(' '.join(lines)))
    r.font.size = Pt(10)
    r.italic = True


def add_table(doc, rows):
    tbl = doc.add_table(rows=0, cols=len(rows[0]))
    tbl.style = 'Table Grid'
    for i, row in enumerate(rows):
        cells = tbl.add_row().cells
        for j, val in enumerate(row):
            if j >= len(cells):
                continue
            cells[j].text = ''
            para = cells[j].paragraphs[0]
            para.paragraph_format.space_after = Pt(2)
            r = para.add_run(clean(val))
            r.font.size = Pt(9)
            if i == 0:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                shade(cells[j]._tc.get_or_add_tcPr(), '0B3D5B')
            elif i % 2 == 0:
                shade(cells[j]._tc.get_or_add_tcPr(), 'F4F7F9')
    doc.add_paragraph()


def add_inline(p, text):
    pattern = r'(\*\*[^*]+\*\*|' + TICK + r'[^' + TICK + r']+' + TICK + r')'
    for token in re.split(pattern, text):
        if not token:
            continue
        if token.startswith('**') and token.endswith('**'):
            r = p.add_run(clean(token[2:-2]))
            r.bold = True
        elif token.startswith(TICK) and token.endswith(TICK):
            r = p.add_run(clean(token[1:-1]))
            r.font.name = 'Consolas'
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0x8B, 0x1A, 0x1A)
        else:
            p.add_run(clean(token))


def render(src, out):
    text = open(src, encoding='utf-8').read()
    meta = {}
    body = []
    for line in text.split('\n'):
        m = re.match(r'^@([A-Z]+):\s*(.*)$', line)
        if m:
            meta[m.group(1)] = m.group(2)
        else:
            body.append(line)

    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.2)
        s.bottom_margin = Cm(2.2)
        s.left_margin = Cm(2.4)
        s.right_margin = Cm(2.0)
    setup_styles(doc)
    cover(doc, meta)
    footer_text(doc, 'LCF - Software Architecture Specification - %s - %s' %
                (meta.get('VOLUME', ''), meta.get('VERSION', '')))

    i = 0
    n = len(body)
    while i < n:
        s = body[i].rstrip()

        if s.startswith(FENCE):
            i += 1
            buf = []
            while i < n and not body[i].startswith(FENCE):
                buf.append(body[i].rstrip('\n'))
                i += 1
            add_code(doc, buf)
            i += 1
            continue

        if s.strip() == '---':
            doc.add_page_break()
            i += 1
            continue

        if s.startswith('> '):
            buf = []
            while i < n and body[i].startswith('> '):
                buf.append(body[i][2:].strip())
                i += 1
            add_note(doc, buf)
            continue

        if s.startswith('|'):
            rows = []
            while i < n and body[i].startswith('|'):
                cells = [c.strip() for c in body[i].strip().strip('|').split('|')]
                if not all(re.fullmatch(r':?-{2,}:?', c) for c in cells):
                    rows.append(cells)
                i += 1
            width = max(len(r) for r in rows)
            rows = [r + [''] * (width - len(r)) for r in rows]
            add_table(doc, rows)
            continue

        m = re.match(r'^(#{1,4})\s+(.*)$', s)
        if m:
            doc.add_heading(clean(m.group(2).strip()), level=len(m.group(1)))
            i += 1
            continue

        m = re.match(r'^(\s*)[-*]\s+(.*)$', s)
        if m:
            lvl = 2 if len(m.group(1)) >= 2 else 1
            p = doc.add_paragraph(style='List Bullet' if lvl == 1 else 'List Bullet 2')
            p.paragraph_format.space_after = Pt(2)
            add_inline(p, m.group(2))
            i += 1
            continue

        m = re.match(r'^\s*\d+[.)]\s+(.*)$', s)
        if m:
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.space_after = Pt(2)
            add_inline(p, m.group(1))
            i += 1
            continue

        if s.strip():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_inline(p, s.strip())
        i += 1

    doc.save(out)
    print('OK ->', out)


if __name__ == '__main__':
    render(sys.argv[1], sys.argv[2])
